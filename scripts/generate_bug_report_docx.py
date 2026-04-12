from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches


ROOT = Path(__file__).resolve().parents[1]
INPUT_PATH = ROOT / "bugs fixing report.txt"
OUTPUT_PATH = ROOT / "SamCard_Bug_Fixing_Report_Professional.docx"


def _set_default_font(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal_style.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SamCard Bug Report and Fixing Summary")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Prepared on {datetime.now():%B %d, %Y}")
    run.italic = True
    run.font.size = Pt(11)

    spacer = document.add_paragraph()
    spacer.add_run(" ")


def _append_content(document: Document, text: str) -> None:
    lines = [line.rstrip() for line in text.splitlines()]

    section_heading_re = re.compile(r"^(\d+)\)\s+(.+)$")
    subsection_heading_re = re.compile(r"^([A-Z])\)\s+(.+)$")
    numbered_re = re.compile(r"^\d+\.\s+")

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            document.add_paragraph("")
            continue

        if set(line) == {"="}:
            continue

        section_match = section_heading_re.match(line)
        if section_match:
            document.add_heading(line, level=1)
            continue

        subsection_match = subsection_heading_re.match(line)
        if subsection_match:
            document.add_heading(line, level=2)
            continue

        if line.endswith(":") and len(line) < 60:
            para = document.add_paragraph()
            para.add_run(line).bold = True
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if numbered_re.match(line):
            body = numbered_re.sub("", line)
            document.add_paragraph(body, style="List Number")
            continue

        document.add_paragraph(line)


def main() -> None:
    if not INPUT_PATH.exists():
        raise FileNotFoundError(f"Input report not found: {INPUT_PATH}")

    source_text = INPUT_PATH.read_text(encoding="utf-8")

    document = Document()
    _set_default_font(document)
    _add_title_block(document)
    _append_content(document, source_text)

    document.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
